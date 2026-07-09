from io import BytesIO

from django.utils import timezone
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .labels import questionnaire_field_label


BRAND_BLUE = '0D416D'
BRAND_RED = 'B8201A'
BRAND_DARK = '172033'
BRAND_MUTED = '64748B'
SOFT_BLUE = 'EAF3FA'
SOFT_RED = 'FFF0EF'
SOFT_GRAY = 'F7FAFC'
LINE = 'D9E4EE'
WHITE = 'FFFFFF'


VALUE_LABELS = {
    'male': 'Мужской',
    'female': 'Женский',
    'school_student': 'Школьник / предварительная заявка',
    'applicant': 'Абитуриент / полная анкета',
    'draft': 'Черновик',
    'submitted': 'Отправлена на проверку',
    'completed': 'Заполнена',
    'approved': 'Принята',
    'rejected': 'Отклонена',
    'updated': 'Обновлена',
    True: 'Да',
    False: 'Нет',
}


SECTIONS = (
    ('Личные данные', ('full_name', 'birth_date', 'gender', 'citizenship', 'marital_status')),
    ('Адрес проживания', ('residence_country', 'residence_region', 'residence_city', 'residence_street', 'residence_house', 'residence_postal_code')),
    ('Паспортные данные', ('passport_number', 'passport_issued_by', 'passport_issue_date', 'passport_expiry_date', 'has_international_passport')),
    ('Контакты', ('phone', 'email', 'extra_phone', 'imo', 'telegram', 'preferred_contact_method')),
    ('Родители / представители', ('parent_full_name', 'parent_relation', 'parent_contacts', 'parent_workplace', 'family_members')),
    ('Образование', ('education_status', 'education_level', 'school_class', 'school_name', 'school_country', 'school_city', 'graduation_year')),
    ('Поступление', ('desired_program', 'admission_goal', 'desired_country', 'desired_city', 'desired_language', 'desired_education_level', 'admission_urgency', 'help_needed')),
    ('Виза', ('has_visa', 'visa_country', 'visa_city', 'visa_valid_until')),
    ('Достижения и языки', ('achievements', 'languages')),
    ('Дополнительная информация', ('hobbies', 'applicant_comment', 'referral_source', 'data_processing_consent')),
)


SCHOOL_STUDENT_FIELDS = {
    'full_name',
    'birth_date',
    'citizenship',
    'residence_country',
    'residence_region',
    'residence_city',
    'phone',
    'email',
    'extra_phone',
    'imo',
    'telegram',
    'parent_full_name',
    'parent_contacts',
    'education_status',
    'school_class',
    'school_name',
    'school_country',
    'school_city',
    'graduation_year',
    'desired_program',
    'desired_country',
    'desired_city',
    'desired_language',
    'applicant_comment',
    'data_processing_consent',
}


def generate_questionnaire_docx(questionnaire):
    document = Document()
    _setup_document(document)
    _add_header(document, questionnaire)
    _add_profile_card(document, questionnaire)
    _add_summary_strip(document, questionnaire)

    section_number = 1
    for title, fields in SECTIONS:
        rows = _section_rows(questionnaire, fields)
        if rows:
            _add_section_card(document, section_number, title, rows)
            section_number += 1

    attachments = list(questionnaire.attachments.all()) if questionnaire.pk else []
    if attachments:
        rows = [
            (attachment.original_name or f'Файл {index}', attachment.file_type or 'Файл')
            for index, attachment in enumerate(attachments, start=1)
        ]
        _add_section_card(document, section_number, 'Вложения анкеты', rows)
        section_number += 1

    _add_manager_notes(document, section_number, questionnaire)
    _add_footer(document)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _setup_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(BRAND_DARK)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(3)


def _add_header(document, questionnaire):
    title = (
        'Предварительная заявка школьника'
        if questionnaire.form_type == questionnaire.FormType.SCHOOL_STUDENT
        else 'Анкета абитуриента'
    )
    generated_at = timezone.localtime(questionnaire.generated_document_at or timezone.now())

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    _set_table_width(table, 9000)
    _remove_table_borders(table)
    table.columns[0].width = Inches(5.2)
    table.columns[1].width = Inches(1.75)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    _set_cell_fill(left, BRAND_BLUE)
    _set_cell_fill(right, BRAND_RED)
    _set_cell_margin(left, 330, 320, 310, 320)
    _set_cell_margin(right, 220, 160, 220, 160)

    _clear_cell(left)
    brand = left.paragraphs[0]
    _add_run(brand, "Student's Life", bold=True, size=21, color=WHITE)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    _add_run(p, title, bold=True, size=18, color=WHITE)
    p = left.add_paragraph()
    _add_run(
        p,
        'Персональное досье для поступления, консультации и сопровождения клиента',
        size=9.5,
        color='DDEAF4',
    )
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    _add_run(p, f'Сформировано: {generated_at:%d.%m.%Y %H:%M}', bold=True, size=8.5, color='DDEAF4')

    _clear_cell(right)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if questionnaire.face_photo:
        try:
            p.add_run().add_picture(questionnaire.face_photo.path, width=Inches(1.22), height=Inches(1.45))
        except (OSError, ValueError):
            _add_photo_placeholder(p)
    else:
        _add_photo_placeholder(p)

    _add_spacer(document, 7)


def _add_profile_card(document, questionnaire):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    _set_table_width(table, 9000)
    _set_table_borders(table, LINE)
    table.columns[0].width = Inches(4.55)
    table.columns[1].width = Inches(2.35)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    _set_cell_fill(left, SOFT_GRAY)
    _set_cell_fill(right, SOFT_RED)
    _set_cell_margin(left, 220, 250, 220, 250)
    _set_cell_margin(right, 220, 220, 220, 220)

    _clear_cell(left)
    p = left.paragraphs[0]
    _add_run(p, 'КЛИЕНТ', bold=True, size=7.8, color=BRAND_MUTED)
    p = left.add_paragraph()
    _add_run(p, questionnaire.full_name or 'ФИО не указано', bold=True, size=17, color=BRAND_BLUE)
    p = left.add_paragraph()
    _add_run(p, _compact_line(questionnaire.phone, questionnaire.email, questionnaire.telegram), size=9.2, color=BRAND_DARK)
    p = left.add_paragraph()
    _add_run(p, _compact_line(questionnaire.desired_program, questionnaire.desired_country, questionnaire.desired_city), size=9.2, color=BRAND_DARK)

    _clear_cell(right)
    p = right.paragraphs[0]
    _add_run(p, 'СТАТУС', bold=True, size=7.8, color=BRAND_MUTED)
    p = right.add_paragraph()
    _add_run(p, questionnaire.get_status_display(), bold=True, size=12, color=BRAND_RED)
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _add_run(p, 'Тип заявки', bold=True, size=7.8, color=BRAND_MUTED)
    p = right.add_paragraph()
    _add_run(p, questionnaire.get_form_type_display(), bold=True, size=9.5, color=BRAND_BLUE)
    _add_spacer(document, 8)


def _add_summary_strip(document, questionnaire):
    items = (
        ('Гражданство', questionnaire.citizenship),
        ('Город', questionnaire.residence_city),
        ('Образование', questionnaire.education_status or questionnaire.education_level),
        ('Год окончания', questionnaire.graduation_year),
    )
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    _set_table_width(table, 9000)
    _remove_table_borders(table)
    for index, (label, value) in enumerate(items):
        cell = table.cell(0, index)
        _set_cell_fill(cell, SOFT_BLUE if index % 2 == 0 else SOFT_GRAY)
        _set_cell_margin(cell, 120, 120, 120, 120)
        _clear_cell(cell)
        p = cell.paragraphs[0]
        _add_run(p, label.upper(), bold=True, size=7.1, color=BRAND_MUTED)
        p = cell.add_paragraph()
        _add_run(p, _format_value(value), bold=True, size=8.6, color=BRAND_BLUE)
    _add_spacer(document, 5)


def _section_rows(questionnaire, fields):
    allowed = SCHOOL_STUDENT_FIELDS if questionnaire.form_type == questionnaire.FormType.SCHOOL_STUDENT else None
    rows = []
    for field in fields:
        if allowed is not None and field not in allowed:
            continue
        value = _format_value(_field_value(questionnaire, field))
        if value != 'Не указано':
            rows.append((questionnaire_field_label(field), value))
    return rows


def _add_section_card(document, number, title, rows):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width(table, 9000)
    _set_table_borders(table, LINE)

    cell = table.cell(0, 0)
    _set_cell_fill(cell, WHITE)
    _set_cell_margin(cell, 170, 210, 160, 210)
    _clear_cell(cell)

    heading = cell.paragraphs[0]
    heading.paragraph_format.space_after = Pt(8)
    _add_run(heading, f'{number:02d}', bold=True, size=8.4, color=BRAND_RED)
    _add_run(heading, f'   {title}', bold=True, size=12.5, color=BRAND_BLUE)

    grid = cell.add_table(rows=0, cols=2)
    grid.autofit = False
    _remove_table_borders(grid)
    for index in range(0, len(rows), 2):
        row = grid.add_row().cells
        _fill_field_cell(row[0], rows[index])
        if index + 1 < len(rows):
            _fill_field_cell(row[1], rows[index + 1])
        else:
            _clear_cell(row[1])
    _add_spacer(document, 7)


def _fill_field_cell(cell, pair):
    label, value = pair
    _set_cell_margin(cell, 70, 80, 85, 80)
    _clear_cell(cell)
    p = cell.paragraphs[0]
    _add_run(p, label.upper(), bold=True, size=7.2, color=BRAND_MUTED)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    _add_run(p, value, size=9.2, color=BRAND_DARK)


def _add_manager_notes(document, number, questionnaire):
    rows = (
        ('Ответственный менеджер', '________________________________'),
        ('Дата проверки', '________________________________'),
        ('Статус анкеты', questionnaire.get_status_display()),
        ('Комментарий', '________________________________'),
        ('Подпись менеджера', '________________________________'),
        ('Подпись клиента / представителя', '________________________________'),
    )
    _add_section_card(document, number, 'Отметки менеджера', rows)


def _add_footer(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    _add_run(
        p,
        "Student's Life • сопровождение поступления • документы • связь с менеджером",
        size=8,
        color=BRAND_MUTED,
    )


def _field_value(questionnaire, field):
    display_method = getattr(questionnaire, f'get_{field}_display', None)
    if callable(display_method):
        value = display_method()
        if value:
            return value
    return getattr(questionnaire, field)


def _format_value(value):
    try:
        if value in VALUE_LABELS:
            return VALUE_LABELS[value]
    except TypeError:
        pass
    if value in (None, '', [], {}):
        return 'Не указано'
    if hasattr(value, 'strftime'):
        if hasattr(value, 'tzinfo') and value.tzinfo:
            return timezone.localtime(value).strftime('%d.%m.%Y %H:%M')
        return value.strftime('%d.%m.%Y')
    if isinstance(value, list):
        items = [_format_list_item(item) for item in value]
        return '\n'.join(f'• {item}' for item in items if item) or 'Не указано'
    if isinstance(value, dict):
        items = [
            f'{questionnaire_field_label(str(key))}: {_format_value(item)}'
            for key, item in value.items()
            if item not in (None, '', [], {})
        ]
        return '\n'.join(items) or 'Не указано'
    return str(value)


def _format_list_item(item):
    if isinstance(item, dict):
        language = item.get('language') or item.get('name') or item.get('title')
        level = item.get('level')
        if language and level:
            return f'{language} — {level}'
        return str(language or item)
    return str(item)


def _compact_line(*values):
    items = [str(value).strip() for value in values if value]
    return ' • '.join(items) if items else 'Не указано'


def _add_photo_placeholder(paragraph):
    _add_run(paragraph, 'ФОТО\n3 x 4', bold=True, size=10, color=WHITE)


def _add_spacer(document, points):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.space_before = Pt(0)


def _add_run(paragraph, text, bold=False, size=9, color=BRAND_DARK):
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def _clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def _set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_cell_margin(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin_name, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{margin_name}'))
        if node is None:
            node = OxmlElement(f'w:{margin_name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def _set_table_width(table, width):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(width))
    tbl_w.set(qn('w:type'), 'dxa')


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'nil')


def _set_table_borders(table, color):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '8')
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)
    for edge in ('insideH', 'insideV'):
        node = borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            borders.append(node)
        node.set(qn('w:val'), 'nil')
